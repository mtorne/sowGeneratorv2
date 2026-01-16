import tempfile
import re
import os
import mimetypes
import base64
from typing import List, Dict, Tuple
from docx import Document
from fastapi import UploadFile
import logging

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for handling document operations"""
    
    @staticmethod
    def extract_placeholders(text: str) -> List[str]:
        """Extract placeholders from text in format {{placeholder}}"""
        return re.findall(r"\{\{(.*?)\}\}", text)
    
    @staticmethod
    def replace_placeholders_in_text(text: str, replacements: Dict[str, str]) -> str:
        """Replace placeholders in text with provided replacements"""
        for key, value in replacements.items():
            text = text.replace(f"{{{{{key}}}}}", value)
        return text
    
    @staticmethod
    async def extract_text_from_docx(file: UploadFile) -> Tuple[str, str]:
        """
        Extract text from uploaded DOCX file
        
        Args:
            file: Uploaded DOCX file
            
        Returns:
            Tuple of (full_text, temp_file_path)
        """
        try:
            contents = await file.read()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(contents)
                tmp_path = tmp.name

            doc = Document(tmp_path)
            
            # Extract text from paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\n"
            
            return full_text, tmp_path
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    @staticmethod
    async def process_diagram(diagram: UploadFile) -> str:
        """
        Process uploaded diagram and return base64 data URI
        
        Args:
            diagram: Uploaded diagram file
            
        Returns:
            Base64 data URI string
        """
        try:
            diagram_bytes = await diagram.read()
            mime_type = mimetypes.guess_type(diagram.filename or "")[0] or "image/png"
            encoded_image = base64.b64encode(diagram_bytes).decode("utf-8")
            return f"data:{mime_type};base64,{encoded_image}"
        except Exception as e:
            logger.error(f"Error processing diagram: {str(e)}")
            raise
    
    @staticmethod
    def cleanup_temp_file(file_path: str) -> None:
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            logger.warning(f"Failed to clean up temp file {file_path}: {str(e)}")


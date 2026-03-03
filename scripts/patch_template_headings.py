#!/usr/bin/env python3
"""One-time script to insert missing implementation-detail subsections into sow_template.docx.

The template already contains a "Implementation Details and Configuration Settings"
Heading 1 paragraph. HIGH AVAILABILITY and MANAGED SERVICES CONFIGURATION (and
optionally BACKUP, DISASTER RECOVERY, etc.) should live as Heading 2 paragraphs
inside that parent section, not as standalone Heading 1 sections.

doc_builder finds headings by keyword substring — once the Heading 2 paragraphs
exist in the template it will inject content in the correct location automatically.

Template structure after patching:
    Heading 1 — Implementation Details and Configuration Settings
    Heading 2 —   High Availability              ← inserted
    Heading 2 —   Managed Services Configuration ← inserted
    ...           (any existing content / OCI Service Sizing follows)

To add more optional subsections (e.g. Backup, Disaster Recovery) in the future,
add entries to SUBSECTIONS below and re-run this script on the template.  The
script is idempotent — headings that already exist are not duplicated.

Usage (run from repo root on the server):
    python scripts/patch_template_headings.py [path/to/sow_template.docx]

The file is modified in-place (a .bak backup is written first).
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document

# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────

# Keyword (lowercase substring) to identify the parent Heading 1 in the template.
PARENT_KEYWORD = "implementation details"

# Subsections to insert as Heading 2 paragraphs inside the parent, IN ORDER.
# Add "Backup", "Disaster Recovery", etc. here if they become in-scope.
SUBSECTIONS: list[str] = [
    "High Availability",
    "Managed Services Configuration",
]


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _is_heading(para) -> bool:
    style = para.style
    while style is not None:
        if style.name.startswith("Heading"):
            return True
        style = getattr(style, "base_style", None)
    return False


def _heading_level(para) -> int | None:
    style = para.style
    while style is not None:
        if style.name.startswith("Heading"):
            try:
                return int(style.name.split()[-1])
            except (ValueError, IndexError):
                return 1
        style = getattr(style, "base_style", None)
    return None


def _make_heading_element(doc: Document, text: str, level: int = 2):
    """Create a detached <w:p> XML element styled as Heading <level>."""
    para = doc.add_paragraph(style=f"Heading {level}")
    para.text = text
    elem = para._element
    elem.getparent().remove(elem)
    return elem


def _all_headings(doc: Document) -> list[tuple[str, int]]:
    """Return list of (text, level) for every heading in body paragraphs."""
    result = []
    for para in doc.paragraphs:
        if _is_heading(para):
            result.append((para.text, _heading_level(para) or 1))
    return result


# ──────────────────────────────────────────────────────────────────────────────
# Main patch logic
# ──────────────────────────────────────────────────────────────────────────────

def patch(template_path: Path) -> None:
    backup_path = template_path.with_suffix(".docx.bak")
    shutil.copy2(template_path, backup_path)
    print(f"Backup written to: {backup_path}")

    doc = Document(str(template_path))

    # ── Idempotency check ────────────────────────────────────────────────────
    existing_texts = {p.text.strip().lower() for p in doc.paragraphs if _is_heading(p)}
    already_present = [s for s in SUBSECTIONS if s.lower() in existing_texts]
    if already_present:
        print(f"Already present (skipping): {already_present}")
    to_insert = [s for s in SUBSECTIONS if s.lower() not in existing_texts]
    if not to_insert:
        print("Nothing to insert — template is already up to date.")
        return

    # ── Find parent heading ──────────────────────────────────────────────────
    parent_para = None
    for para in doc.paragraphs:
        if _is_heading(para) and PARENT_KEYWORD in para.text.lower():
            parent_para = para
            break

    if parent_para is None:
        print(f"ERROR: Could not find a heading containing '{PARENT_KEYWORD}'.")
        print("Available headings:")
        for text, lvl in _all_headings(doc):
            print(f"  [H{lvl}] {text!r}")
        sys.exit(1)

    print(f"Parent heading found: [{_heading_level(parent_para)}] {parent_para.text!r}")

    # ── Find the insert-before anchor ────────────────────────────────────────
    # We want to insert the new Heading 2 paragraphs immediately after the
    # parent heading (before any existing next sibling heading).
    #
    # Strategy: walk doc.paragraphs from parent_para onwards and find the
    # first heading at level ≤ parent level that is NOT the parent itself.
    # New subsections go in before that next peer heading (or at end of section).
    parent_level = _heading_level(parent_para) or 1
    paragraphs = list(doc.paragraphs)
    parent_idx = paragraphs.index(parent_para)

    anchor_elem = None  # insert-before this element; None means insert after parent
    for para in paragraphs[parent_idx + 1:]:
        if _is_heading(para):
            lvl = _heading_level(para) or 1
            if lvl <= parent_level:
                # Next heading at same or higher level — insert before this
                anchor_elem = para._element
                print(f"Inserting before: [{lvl}] {para.text!r}")
                break
            # Heading 2 inside the section — insert before existing sub-headings
            # so our new ones come first (only if they should precede them).
            # For now we insert immediately after parent — keep existing order.

    # ── Insert subsections ───────────────────────────────────────────────────
    if anchor_elem is not None:
        # Insert in reverse order so final order matches SUBSECTIONS list
        for text in reversed(to_insert):
            elem = _make_heading_element(doc, text, level=parent_level + 1)
            anchor_elem.addprevious(elem)
            print(f"  Inserted [H{parent_level + 1}] '{text}' (before anchor)")
    else:
        # No next peer heading found — append after parent element in reverse
        insert_after = parent_para._element
        for text in reversed(to_insert):
            elem = _make_heading_element(doc, text, level=parent_level + 1)
            insert_after.addnext(elem)
            print(f"  Inserted [H{parent_level + 1}] '{text}' (after parent)")

    doc.save(str(template_path))
    print(f"\nTemplate patched and saved: {template_path}")

    # ── Verification ─────────────────────────────────────────────────────────
    print("\nVerification — headings around the patched area:")
    all_h = _all_headings(Document(str(template_path)))
    for i, (text, lvl) in enumerate(all_h):
        if PARENT_KEYWORD in text.lower():
            start = max(0, i - 1)
            end = min(len(all_h), i + len(to_insert) + 3)
            for t, l in all_h[start:end]:
                marker = "  ← NEW" if t in to_insert else ""
                print(f"  [H{l}] {t!r}{marker}")
            break


# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
    else:
        # Default: repo-relative path used in production
        path = Path(__file__).resolve().parent.parent / "app" / "sow_template.docx"

    if not path.exists():
        print(f"ERROR: Template not found at {path}")
        sys.exit(1)

    patch(path)

"""Document assembly services."""

from __future__ import annotations

import datetime
import json as _json
import logging
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from app.services.doc_style_constants import (
    BODY_COLOR,
    BOM_HEADER_FILL,
    CELL_MARGIN_BOTTOM,
    CELL_MARGIN_LEFT,
    CELL_MARGIN_RIGHT,
    CELL_MARGIN_TOP,
    CUSTOMER_COLOR,
    INNER_BORDER_COLOR,
    INNER_BORDER_SIZE,
    ORACLE_COLOR,
    OUTER_BORDER_COLOR,
    OUTER_BORDER_SIZE,
    PENDING_COLOR,
    STYLE_HEADING3,
    STYLE_LIST_BULLET,
    STYLE_LIST_BULLET2,
    STYLE_LIST_BULLET3,
    STYLE_LIST_NUMBERED,
    TABLE_ALT_FILL_DARK,
    TABLE_ALT_FILL_LIGHT,
    TABLE_HEADER_FILL,
    TABLE_HEADER_TEXT,
    TABLE_STYLE_BOM,
)

logger = logging.getLogger(__name__)

# Maps canonical section names → keyword to match in template heading text
SECTION_HEADING_KEYWORDS: dict[str, str] = {
    "SOW VERSION HISTORY":                          "version history",
    "STATUS AND NEXT STEPS":                        "status and next",
    "PROJECT PARTICIPANTS":                         "project participants",
    "COMPANY PROFILE":                              "company profile",
    "IN SCOPE APPLICATION":                         "in scope application",
    "PROJECT OVERVIEW":                             "project overview",
    # "Scope" H3 lives inside Project Overview H1.
    # Prefix "=" means exact-match (avoids matching "In Scope Application").
    "SCOPE":                                        "=scope",
    "CURRENT STATE ARCHITECTURE":                   "current state architecture",
    # H3 sub-section "Current State Architecture - Description"
    "CURRENT STATE ARCHITECTURE DESCRIPTION":       "current state architecture - description",
    "CURRENTLY USED TECHNOLOGY STACK":              "currently used technology",
    "OCI SERVICE SIZING AND AMOUNTS":               "oci service sizing",
    "FUTURE STATE ARCHITECTURE":                    "future state architecture",
    "ARCHITECTURE DEPLOYMENT OVERVIEW":             "architecture deployment",
    "ARCHITECTURE COMPONENTS":                      "architecture components",
    "IMPLEMENTATION DETAILS":                       "implementation details",
    # Template uses "Major Project Milestones" — match on that substring.
    "MILESTONE PLAN":                               "major project milestones",
    "SECURITY":                                     "security",
    "HIGH AVAILABILITY":                            "high availability",
    "BACKUP STRATEGY":                              "backup strategy",
    "DISASTER RECOVERY":                            "disaster recovery",
    "MANAGED SERVICES CONFIGURATION":               "managed services",
    # Architect review — appended at end when not found in template heading
    "ARCHITECT REVIEW":                             "architect review",
    "CLOSING FEEDBACK":                             "closing feedback",
}

# When a section's own heading is absent from the template, insert it immediately
# after the last paragraph of this sibling section rather than appending at the end.
# Processed in order — DR anchors to BACKUP STRATEGY which itself anchors to HA.
# ARCHITECT REVIEW anchors to DISASTER RECOVERY (itself dynamically injected),
# placing it after DR and before OCI SERVICE SIZING AND AMOUNTS.
_SECTION_INSERT_AFTER: dict[str, str] = {
    "BACKUP STRATEGY":   "high availability",
    "DISASTER RECOVERY": "backup strategy",
    "ARCHITECT REVIEW":  "disaster recovery",
}

# Sections whose template body content should be fully cleared before LLM injection.
# Use for sections that are 100% LLM-generated with no useful template intro text.
_FULL_CLEAR_SECTIONS = frozenset({
    "ARCHITECTURE COMPONENTS",
    # NOTE: SCOPE is intentionally NOT here — the template colored-box headers
    # ("Initial understanding of the scope", "Desired Outcome of customer",
    # "Desired outcome agreed with Oracle") must survive so users can fill them.
    # The Description H3 sub-section has only generic placeholder intro text.
    "CURRENT STATE ARCHITECTURE DESCRIPTION",
    # STATUS AND NEXT STEPS has only generic template intro text ("Current status
    # and what needs to happen next…") — replace it entirely with the LLM output.
    "STATUS AND NEXT STEPS",
    # OCI SERVICE SIZING AND AMOUNTS — the template has several generic intro
    # paragraphs; replace them all with the single LLM-generated sentence.
    "OCI SERVICE SIZING AND AMOUNTS",
    # Operative sections — fully LLM-generated from diagram analysis
    "MILESTONE PLAN",
    "HIGH AVAILABILITY",
    "BACKUP STRATEGY",
    "DISASTER RECOVERY",
    # SECURITY and MANAGED SERVICES CONFIGURATION are now LLM-generated from
    # the target architecture analysis; clear template boilerplate before injection.
    "SECURITY",
    "MANAGED SERVICES CONFIGURATION",
})

# Sections where LLM content should be injected RIGHT AFTER the heading,
# BEFORE any surviving template paragraphs.
# SCOPE was removed from this set — it now uses _inject_scope_boxes() which
# removes only LLM-generated intro text and injects content into the colored
# scope-box paragraphs via soft breaks, leaving template text intact.
_INJECT_AT_TOP_SECTIONS: frozenset[str] = frozenset()

# Sections whose LLM output uses sub-topic labels (Networking, Security, Compute …)
# and should be formatted with bold labels + sentence-level bullet points.
_LABELED_FORMAT_SECTIONS = frozenset({
    "IMPLEMENTATION DETAILS",
    # Architect review uses the same label+bullet format for its sub-sections.
    "ARCHITECT REVIEW",
    # Status and next steps uses Completed / Pending labels + bullet lines.
    "STATUS AND NEXT STEPS",
    # Operative sections — Phase N / sub-topic label format
    "MILESTONE PLAN",
    "HIGH AVAILABILITY",
    "BACKUP STRATEGY",
    "DISASTER RECOVERY",
})

# Sections where the writer emits JSON instead of prose.
# doc_builder tries to parse the content as JSON and call _inject_structured_section;
# on parse failure it gracefully falls back to _inject_formatted_blocks_after_element.
_STRUCTURED_OUTPUT_SECTIONS: frozenset[str] = frozenset({
    "MILESTONE PLAN",
    "HIGH AVAILABILITY",
    "BACKUP STRATEGY",
    "DISASTER RECOVERY",
})

# Ordered (field_name, display_label) pairs used by _inject_structured_section
# for the HA / BACKUP / DR sections.  MILESTONE PLAN uses its own rendering path.
_STRUCTURED_SECTION_LABELS: dict[str, list[tuple[str, str]]] = {
    "HIGH AVAILABILITY": [
        ("oci_ha_capabilities",     "OCI HA Capabilities"),
        ("redundancy_architecture", "Redundancy Architecture"),
        ("failover_strategy",       "Failover Strategy"),
        ("rto_rpo_targets",         "RTO/RPO Targets"),
    ],
    "BACKUP STRATEGY": [
        ("data_backup",         "Data Backup"),
        ("application_backup",  "Application Backup"),
        ("recovery_procedures", "Recovery Procedures"),
        ("retention_policy",    "Retention Policy"),
    ],
    "DISASTER RECOVERY": [
        ("dr_strategy",           "DR Strategy"),
        ("geographic_redundancy", "Geographic Redundancy"),
        ("data_replication",      "Data Replication"),
        ("dr_testing_plan",       "DR Testing Plan"),
    ],
}

# Sections whose LLM output uses hierarchical bullet lines (L1/L2/L3 indentation).
# Each newline-separated line becomes its own paragraph; leading spaces determine
# the indent level: 0 → L1 (0 twips), 2 spaces → L2 (360 twips), 4+ → L3 (720 twips).
#
# NOTE: CURRENT STATE ARCHITECTURE DESCRIPTION was removed — the reference SoW
# uses a single flowing prose paragraph, not hierarchical bullets.  The prompt
# was updated to request narrative prose; the plain paragraph renderer is used.
_HIERARCHICAL_BULLET_SECTIONS = frozenset({
    # Currently empty — kept for future sections that need indented bullet rendering.
})

# Sections where the LLM is instructed to produce a SINGLE introductory sentence.
# Even when the model emits extra paragraphs, only the first one is injected.
# This prevents LLM verbosity from adding unwanted paragraphs above tables.
_SINGLE_SENTENCE_SECTIONS = frozenset({
    "OCI SERVICE SIZING AND AMOUNTS",
    "PROJECT PARTICIPANTS",
    "CURRENTLY USED TECHNOLOGY STACK",
    # IN SCOPE APPLICATION is table-driven; the LLM writes only a minimal lead-in.
    "IN SCOPE APPLICATION",
    # FUTURE STATE ARCHITECTURE is a container heading whose sub-sections
    # (Target Architecture Diagram, Architecture Deployment Overview, etc.)
    # carry all the detail.  The reference SoW has zero body text between the
    # Heading 1 and the first Heading 3, so we limit injection to at most one
    # short introductory sentence to avoid polluting the layout.
    "FUTURE STATE ARCHITECTURE",
})

# Known sub-topic label set for LABELED_FORMAT_SECTIONS
_SUB_TOPIC_LABELS = frozenset({
    "networking",
    "security",
    "compute",
    "storage and databases",
    "devops and management",
    "storage",
    "databases",
    # Architect Review sub-sections
    "generation quality",
    "data gaps",
    "next steps",
    "recommendations",
    # Status and Next Steps sub-sections
    "completed",
    "pending",
    # (Milestone Plan phase labels removed — milestones are now appended
    # to the existing template table, not rendered as labelled blocks.)
    # High Availability sub-sections
    "oci ha capabilities",
    "redundancy architecture",
    "failover strategy",
    "rto/rpo targets",
    # Backup Strategy sub-sections
    "data backup",
    "application backup",
    "recovery procedures",
    "retention policy",
    # Disaster Recovery sub-sections
    "dr strategy",
    "geographic redundancy",
    "data replication",
    "dr testing plan",
})

# Regex to split prose into sentences at ". " followed by an uppercase letter
_SENTENCE_SPLIT_RE = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')

# Literal string placed by LLM for unknown values — rendered red in the DOCX.
_PENDING_MARKER = "PENDING TO REVIEW"

# Known capitalisation errors baked into the sow_template.docx headings.
# Applied to every <w:t> element in the document so both headings and their
# cached TOC counterparts are corrected on every build.
_HEADING_TEXT_FIXES: tuple[tuple[str, str], ...] = (
    ("NExt STEPS",    "Next Steps"),
    (" STate ",       " State "),
    ("CuRrently",     "Currently"),   # may also appear split across runs — see below
)

# Known typos and informal notations in the template body text.
_BODY_TEXT_FIXES: tuple[tuple[str, str], ...] = (
    ("Capability/Metrci",           "Capability/Metric"),
    ("Possible arias to include",   "Possible areas to include"),
    ("NB**  ",                      "Note: "),
    ("NB** ",                       "Note: "),
)

# Placeholder text patterns to remove when found inside a section.
# NOTE: "Initial understanding of the scope", "Desired Outcome, as jointly agreed",
# and "Any change in the objectives and scope" have been deliberately removed so
# the template's colored scope-box headers survive injection.
_PLACEHOLDER_RE = re.compile(
    r"<add information here>"
    r"|<Information to be filled in[^>]*>"
    r"|<Diagram to be provided[^>]*>"
    r"|<Elaborate on the architecture[^>]*>"
    r"|<Enumerate the target[^>]*>"
    r"|<Final thoughts[^>]*>"
    r"|\[Add description here\]"
    r"|\[add information here\]"
    r"|Relevant aspects of the architecture, how to deploy.*$"
    r"|Desired project outcome is to be provided by"
    r"|If the Statement of Work is filled in separately",
    re.IGNORECASE,
)

# Heading title suffixes (lower-case, without leading customer name) that mark
# headings whose FIRST run is a blank placeholder for the customer name.
_CUSTOMER_HEADING_PREFIXES: frozenset[str] = frozenset({
    "company profile",
    "offboarding",
})

# Known prefixes in template runs that are immediately followed by a blank
# customer-name run.  Ordered longest-match first to avoid short-prefix false
# positives.
_CUSTOMER_PREFIX_SUFFIXES = (
    "the Non-Disclosure Agreement between ",
    "agreement between ",           # "require mutual agreement between [customer]"
    "free of charge for the ",
    "to enable ",
    "available to ",
    "Oracle-",
    "Oracle and ",              # "shared between Oracle and [customer]"
    "initial format by ",
    "granted for ",
    "the  Oracle",              # "the [customer] Oracle Labs" — blank INSIDE run
    "under the  Oracle",
    "based on the ",
    "falls under ",
    "fall under ",
    "ACE/Sales and ",
    "for the ",
    "provided by ",
    "under ",
    "The ",
)


def _build_para_elem(
    text: str, bold: bool = False, list_style: bool = False
) -> object:
    """Build a ``<w:p>`` XML element containing one or more runs.

    If *text* contains the :data:`_PENDING_MARKER` literal, the marker
    segments are emitted as separate red-coloured runs so they appear in
    red in the rendered DOCX.  All other text is emitted in the normal
    (or bold) colour.

    When *list_style* is ``True`` a ``<w:pPr><w:pStyle w:val="ListParagraph"/>``
    element is prepended so the paragraph picks up the template's List Paragraph
    style (indentation, spacing) matching the reference document convention.

    For proper Word list styling (with real bullets from template numbering),
    use :func:`_build_word_style_para_elem` instead.
    """
    new_para = OxmlElement("w:p")

    if list_style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "ListParagraph")
        pPr.append(pStyle)
        new_para.insert(0, pPr)

    def _add_run(t: str, red: bool = False) -> None:
        if not t:
            return
        r = OxmlElement("w:r")
        if bold or red:
            rPr = OxmlElement("w:rPr")
            if bold:
                rPr.append(OxmlElement("w:b"))
            if red:
                c = OxmlElement("w:color")
                c.set(qn("w:val"), PENDING_COLOR)
                rPr.append(c)
            r.append(rPr)
        t_elem = OxmlElement("w:t")
        t_elem.text = t
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t_elem)
        new_para.append(r)

    if _PENDING_MARKER in text:
        parts = text.split(_PENDING_MARKER)
        for idx, part in enumerate(parts):
            _add_run(part, red=False)
            if idx < len(parts) - 1:
                _add_run(_PENDING_MARKER, red=True)
    else:
        _add_run(text, red=False)

    return new_para


def _build_word_style_para_elem(text: str, style_id: str) -> object:
    """Build a ``<w:p>`` that uses a named Word paragraph style from the template.

    Unlike :func:`_build_para_elem` (which sets ListParagraph with a text-dash
    prefix), this function sets the paragraph style to a proper Word list style
    (e.g. ``NormalBodyBullet1``, ``NormalBodyBullet2``, ``NumberedList1``,
    ``Heading3``) so that the bullet character, indentation, and spacing all
    derive from the template's own style definitions extracted from the
    reference document.

    ``PENDING TO REVIEW`` markers are rendered in red, same as _build_para_elem.

    Strip any leading ``– `` or ``- `` dash prefix from *text* before calling
    this function — the Word list style supplies the bullet character.
    """
    new_para = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    new_para.insert(0, pPr)

    def _add_run(t: str, red: bool = False) -> None:
        if not t:
            return
        r = OxmlElement("w:r")
        if red:
            rPr = OxmlElement("w:rPr")
            c = OxmlElement("w:color")
            c.set(qn("w:val"), PENDING_COLOR)
            rPr.append(c)
            r.append(rPr)
        t_elem = OxmlElement("w:t")
        t_elem.text = t
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t_elem)
        new_para.append(r)

    if _PENDING_MARKER in text:
        parts = text.split(_PENDING_MARKER)
        for idx, part in enumerate(parts):
            _add_run(part, red=False)
            if idx < len(parts) - 1:
                _add_run(_PENDING_MARKER, red=True)
    else:
        _add_run(text, red=False)

    return new_para


def _strip_bullet_prefix(text: str) -> str:
    """Remove a leading ``– ``, ``- ``, or ``• `` bullet prefix from *text*.

    Word list styles supply their own bullet character; text content must not
    carry a duplicate manual prefix.
    """
    for prefix in ("– ", "- ", "• ", "* "):
        if text.startswith(prefix):
            return text[len(prefix):]
    return text


def _build_indented_para_elem(text: str, left_twips: int = 0) -> object:
    """Build a ``<w:p>`` with optional left indentation.

    Like :func:`_build_para_elem` but supports hierarchical indentation for
    bullet lists.  The *left_twips* value is applied as ``<w:ind w:left>``.
    ``PENDING TO REVIEW`` markers are rendered in red, same as _build_para_elem.
    """
    new_para = OxmlElement("w:p")

    if left_twips:
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left_twips))
        pPr.append(ind)
        new_para.insert(0, pPr)

    def _add_run(t: str, red: bool = False) -> None:
        if not t:
            return
        r = OxmlElement("w:r")
        if red:
            rPr = OxmlElement("w:rPr")
            c = OxmlElement("w:color")
            c.set(qn("w:val"), "FF0000")
            rPr.append(c)
            r.append(rPr)
        t_elem = OxmlElement("w:t")
        t_elem.text = t
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t_elem)
        new_para.append(r)

    if _PENDING_MARKER in text:
        parts = text.split(_PENDING_MARKER)
        for idx, part in enumerate(parts):
            _add_run(part, red=False)
            if idx < len(parts) - 1:
                _add_run(_PENDING_MARKER, red=True)
    else:
        _add_run(text)

    return new_para


class DocumentBuilder:
    """Injects generated text into a DOCX template."""

    def __init__(
        self,
        template_path: Path,
        customer_name: str = "",
        project_name: str = "",
    ) -> None:
        self.template_path = template_path
        self.customer_name = customer_name.strip()
        self.project_name = project_name.strip()

    def _load_or_create_template(self) -> Document:
        if self.template_path.exists():
            return Document(str(self.template_path))
        logger.warning("Template not found at %s. Using generated fallback template", self.template_path)
        return Document()

    # ------------------------------------------------------------------
    # Customer / project name substitution
    # ------------------------------------------------------------------

    def _substitute_names(self, doc: Document) -> None:
        """Replace 'Customer1' / 'Project1' placeholders with actual names.

        Operates at the XML level so that ALL paragraphs are covered,
        including those inside text boxes (w:txbxContent) and table cells
        which are not surfaced by doc.paragraphs.

        Three passes per paragraph:
        1. Replace literal placeholder tokens ('Customer1', 'Project1') in
           every run's <w:t> text, with an anti-duplication check against
           the following run.
        2. Fill blank/space-only runs that sit between context runs where
           the customer name belongs (identified by the preceding run's
           known suffix).  Includes a guard against filling when the next
           run already starts with the customer name.
        3. Collapse consecutive runs whose text equals the customer name
           (happens when the template had both a 'Customer1' token and an
           adjacent already-filled run).
        4. Blank run at paragraph start (i==0): fill when the paragraph has
           other non-empty content and does not already contain the name.
        """
        if not self.customer_name and not self.project_name:
            return

        token_map: dict[str, str] = {}
        if self.customer_name:
            token_map["Customer1"] = self.customer_name
        if self.project_name:
            token_map["Project1"] = self.project_name

        def _process_p_element(p_elem) -> None:
            # Build an ordered run list: direct-child <w:r> elements and
            # <w:fldSimple> children interleaved in document order.
            #
            # Word templates often store Customer1 / Project1 as a DOCPROPERTY
            # field:  <w:fldSimple w:instr="DOCPROPERTY 01_Customer">
            #           <w:r><w:t> </w:t></w:r>   ← grandchild placeholder
            #         </w:fldSimple>
            # A plain findall("w:r") returns only DIRECT children and misses
            # these grandchild runs.  When the blank placeholder sits inside a
            # fldSimple, Pass 2's context-suffix matching never sees it and the
            # customer name is never injected.
            #
            # Iterating p_elem's direct children in order and expanding
            # fldSimple elements preserves the correct left-context (the run
            # BEFORE the fldSimple still appears at index i-1 so suffix
            # matching works correctly).
            r_elems = []
            for child in p_elem:
                if child.tag == qn("w:r"):
                    r_elems.append(child)
                elif child.tag == qn("w:fldSimple"):
                    r_elems.extend(child.findall(qn("w:r")))
                elif child.tag == qn("w:hyperlink"):
                    # TOC entries wrap their text runs inside <w:hyperlink>.
                    # Including those runs here ensures Customer1 / Project1
                    # tokens inside cached TOC text are also substituted.
                    r_elems.extend(child.findall(qn("w:r")))

            all_r_for_pass1 = r_elems  # fldSimple runs now included

            if not all_r_for_pass1:
                return

            def _get_t(r) -> object | None:
                return r.find(qn("w:t"))

            def _text(r) -> str:
                t = _get_t(r)
                return (t.text or "") if t is not None else ""

            def _set_text(r, value: str) -> None:
                t = _get_t(r)
                if t is not None:
                    t.text = value
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                else:
                    new_t = OxmlElement("w:t")
                    new_t.text = value
                    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    r.append(new_t)

            # Pass 1: literal token replacement (direct runs + fldSimple field runs)
            for idx, r in enumerate(r_elems):
                for token, replacement in token_map.items():
                    current = _text(r)
                    if token in current:
                        new_val = current.replace(token, replacement)
                        # Anti-duplication: if next run already starts with
                        # the replacement we just inserted, strip it there.
                        if idx + 1 < len(r_elems):
                            next_text = _text(r_elems[idx + 1])
                            if next_text.startswith(replacement):
                                _set_text(r_elems[idx + 1], next_text[len(replacement):])
                        _set_text(r, new_val)
            # (fldSimple runs are now part of r_elems and processed above)

            if not r_elems:
                return  # no direct-child runs → skip Pass 2/3

            # Refresh run texts after pass 1
            run_texts = [_text(r) for r in r_elems]
            para_text = "".join(run_texts)

            if not self.customer_name:
                return

            # Pass 2: fill blank/space-only runs by context suffix
            # Only fills a blank run when the PRECEDING run ends with one of
            # the known prefix strings — this avoids false positives on blank
            # runs inside headings, table cells, and other non-name slots.
            for i, r in enumerate(r_elems):
                txt = run_texts[i]
                if txt.strip() != "":
                    continue  # not a blank run

                if i == 0:
                    # Special case: leading blank run in a heading whose title
                    # (formed by joining all SUBSEQUENT runs) is a known
                    # "customer-prefixed" heading (e.g. " Company Profile").
                    # Fill it with the customer name only when the paragraph
                    # text does not already contain the customer name.
                    if not self.customer_name:
                        continue
                    if self.customer_name in para_text:
                        continue  # already substituted — skip
                    rest_text = "".join(run_texts[1:]).strip().lower()
                    if rest_text in _CUSTOMER_HEADING_PREFIXES:
                        # Preserve a space separator between customer name and
                        # the heading title that follows in the next run.
                        fill = self.customer_name
                        if (
                            i + 1 < len(r_elems)
                            and not run_texts[i + 1].startswith(" ")
                        ):
                            fill = self.customer_name + " "
                        _set_text(r, fill)
                        run_texts[i] = fill
                        logger.debug(
                            "doc_builder.customer_name_injected_heading rest=%r", rest_text
                        )
                    continue  # no preceding context for suffix-based fill

                # Guard: if the following run already starts with the
                # customer name, this blank run is adjacent to the real
                # value — do not fill to avoid duplication.
                if i + 1 < len(r_elems) and run_texts[i + 1].startswith(self.customer_name):
                    continue

                prev_text = run_texts[i - 1]
                for suffix in _CUSTOMER_PREFIX_SUFFIXES:
                    if prev_text.endswith(suffix):
                        _set_text(r, self.customer_name)
                        run_texts[i] = self.customer_name
                        logger.debug(
                            "doc_builder.customer_name_injected run_idx=%d prev_suffix=%r",
                            i, suffix,
                        )
                        break

            # Pass 3: collapse consecutive customer-name runs
            # (template had two adjacent placeholder runs for the same slot)
            last_name_idx: int | None = None
            run_texts = [_text(r) for r in r_elems]  # refresh
            for i, r in enumerate(r_elems):
                txt = run_texts[i]
                if txt == self.customer_name:
                    if last_name_idx is not None:
                        _set_text(r, "")  # clear the duplicate
                        logger.debug("doc_builder.customer_name_deduped run_idx=%d", i)
                    else:
                        last_name_idx = i
                elif txt.strip():
                    last_name_idx = None  # reset on non-empty, non-name text

        # Iterate over EVERY <w:p> in the document, including those inside
        # text boxes (w:txbxContent), headers, footers, and table cells.
        for p_elem in doc.element.body.iter(qn("w:p")):
            try:
                _process_p_element(p_elem)
            except Exception:
                logger.debug("doc_builder.substitute_names_para_error", exc_info=True)

        logger.info(
            "doc_builder.names_substituted customer=%r project=%r",
            self.customer_name,
            self.project_name,
        )

    def _normalize_template_artifacts(self, doc: Document) -> None:
        """Fix known typos, capitalisation errors, and artefacts in the template.

        Runs after :meth:`_substitute_names` on every build so that the output
        DOCX is clean regardless of whether the source sow_template.docx has
        been patched.  Covers:

        * Heading capitalisation errors (``NExt STEPS``, ``CuRrently``, ``STate``)
          — applied to every ``<w:t>`` so both headings and their cached TOC
          entries are corrected in a single pass.
        * Body-text typos (``Metrci``, ``arias``, ``NB**`` notation).
        * Removal of the "Dropdown Options" form-control (SDT) on the cover page.
        """
        all_fixes = _HEADING_TEXT_FIXES + _BODY_TEXT_FIXES
        body = doc.element.body

        # Pass 1: simple per-<w:t> replacements
        for t_el in body.iter(qn("w:t")):
            if not t_el.text:
                continue
            for old, new in all_fixes:
                if old in t_el.text:
                    t_el.text = t_el.text.replace(old, new)

        # Pass 2: "CuRrently" split-run case — template stores it as three
        # consecutive runs: "Cu" | "R" | "rently …".  Detect by checking
        # neighbours within the same paragraph.
        for p_elem in body.iter(qn("w:p")):
            wts = list(p_elem.iter(qn("w:t")))
            for i, t_el in enumerate(wts):
                if t_el.text != "R":
                    continue
                prev = wts[i - 1].text if i > 0 else ""
                nxt = wts[i + 1].text if i + 1 < len(wts) else ""
                if (prev or "").endswith("Cu") and (nxt or "").startswith("rently"):
                    t_el.text = "r"

        # Pass 3: remove "Dropdown Options" structured-document-tags (SDTs).
        # These are form-control artefacts on the cover page that render as
        # visible text in some Word versions.
        for sdt in list(body.iter(qn("w:sdt"))):
            sdt_text = "".join(t.text or "" for t in sdt.iter(qn("w:t")))
            if "Dropdown" in sdt_text:
                parent = sdt.getparent()
                if parent is not None:
                    parent.remove(sdt)

        logger.debug("doc_builder.template_artifacts_normalized")

    def build(
        self,
        sections: list[tuple[str, str]],
        output_dir: Path,
        diagram_images: dict[str, list[tuple[str, bytes]]] | dict[str, bytes] | None = None,
        project_context: dict | None = None,
        include_architect_review: bool = False,
        excluded_sections: frozenset[str] | None = None,
    ) -> str:
        """Inject sections into template headings and save DOCX.

        Args:
            sections: List of (section_name, content) tuples in canonical order.
            output_dir: Directory where the output DOCX is written.
            diagram_images: Optional mapping of ``"current"`` / ``"target"`` to
                either a list of ``(filename, bytes)`` tuples (multiple diagrams)
                or a single raw ``bytes`` value (legacy single-image format).
                All images for each role are embedded in the document with
                descriptive captions.
            project_context: Full project context dict from the API request (client,
                project_name, scope, industry, services …).  Used to populate
                Company Profile, In Scope Application, and Acceptance Criteria tables.
            include_architect_review: When ``False`` (default) the ``ARCHITECT
                REVIEW`` section is stripped from *sections* before injection.
                Set to ``True`` to retain it as an internal audit trail — it
                should never be present in a customer-facing deliverable.
            excluded_sections: Optional set of uppercased section names to omit
                from DOCX injection (e.g. ``{"HIGH AVAILABILITY", "DISASTER RECOVERY"}``).
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        output_name = f"output_{uuid4().hex}.docx"
        output_path = output_dir / output_name

        _excluded = excluded_sections or frozenset()
        if not include_architect_review:
            _excluded = _excluded | {"ARCHITECT REVIEW"}
        if _excluded:
            sections = [
                (name, content)
                for name, content in sections
                if name not in _excluded
            ]

        doc = self._load_or_create_template()
        self._substitute_names(doc)
        self._normalize_template_artifacts(doc)
        self._fill_project_tables(doc, project_context=project_context)

        # Physically remove excluded section headings + their template body content so
        # they don't appear as empty stubs in the final DOCX.
        if _excluded and self.template_path.exists():
            for excl_name in _excluded:
                self._delete_section_from_template(doc, excl_name)

        if self.template_path.exists():
            for section_name, content in sections:
                if not content or not content.strip():
                    continue
                injected = self._inject_section(doc, section_name, content)
                if not injected:
                    anchor = _SECTION_INSERT_AFTER.get(section_name.upper())
                    if anchor:
                        injected = self._inject_after_known_section(
                            doc, section_name, content, after_keyword=anchor
                        )
                    if not injected:
                        logger.warning("doc_builder.section_not_found section=%s — appending at end", section_name)
                        self._append_section(doc, section_name, content)
        else:
            # Fallback: build from scratch
            for section_name, content in sections:
                doc.add_heading(section_name.title(), level=1)
                for block in content.split("\n\n"):
                    if block.strip():
                        doc.add_paragraph(block.strip())

        if diagram_images:
            self._insert_diagram_images(doc, diagram_images)

        self._apply_heading1_page_breaks(doc)
        self._suppress_blank_heading_page_breaks(doc)
        doc.save(str(output_path))
        logger.info("Saved generated SoW document: %s", output_path)
        return output_name

    def build_markdown(self, full_document: str, output_dir: Path) -> str:
        """Save generated content as markdown and return file name."""
        output_dir.mkdir(parents=True, exist_ok=True)
        output_name = f"output_{uuid4().hex}.md"
        output_path = output_dir / output_name
        output_path.write_text(full_document, encoding="utf-8")
        logger.info("Saved generated SoW markdown: %s", output_path)
        return output_name

    # ------------------------------------------------------------------
    # Diagram image insertion  (Feature B)
    # ------------------------------------------------------------------

    def _insert_diagram_images(
        self,
        doc: Document,
        diagram_images: dict[str, list[tuple[str, bytes]]] | dict[str, bytes],
    ) -> None:
        """Replace placeholder images with the actual uploaded architecture diagrams.

        Searches for the headings "Current State Architecture - Diagram" and
        "Target Architecture Diagram" in the document, finds the placeholder
        drawing paragraph that follows each heading (within 10 paragraphs), and
        inserts all uploaded diagrams for that role with descriptive captions.

        When multiple images are uploaded for a role, all are embedded sequentially.
        A caption paragraph ("Figure N: …") is added below each image.

        Args:
            diagram_images: Mapping of ``"current"`` or ``"target"`` to either:
                * a list of ``(filename, bytes)`` tuples — multiple images, or
                * a raw ``bytes`` value — single image (legacy format, no caption).
        """
        _DRAWING_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _DRAWING_TAG = f"{{{_DRAWING_NS}}}drawing"

        # Normalise to list[tuple[str, bytes]] regardless of input format
        def _normalise(value) -> list[tuple[str, bytes]]:
            if isinstance(value, (bytes, bytearray)):
                return [("diagram.png", bytes(value))]
            if isinstance(value, list):
                return [(str(name), bytes(data)) for name, data in value]
            return []

        slot_map = [
            ("current", "current state architecture - diagram"),
            ("target", "target architecture diagram"),
        ]

        all_paras = doc.paragraphs

        for role, keyword in slot_map:
            raw = diagram_images.get(role)
            if not raw:
                logger.debug("doc_builder.diagram_skip role=%s (no image bytes)", role)
                continue
            images = _normalise(raw)
            if not images:
                continue

            # Find the matching heading paragraph
            heading_idx: int | None = None
            for i, p in enumerate(all_paras):
                if keyword in p.text.lower() and self._is_heading_style(p):
                    heading_idx = i
                    break

            if heading_idx is None:
                logger.warning(
                    "doc_builder.diagram_heading_not_found role=%s keyword=%r", role, keyword
                )
                continue

            # Find the first placeholder drawing paragraph after the heading
            placeholder_para = None
            search_limit = min(heading_idx + 10, len(all_paras))
            for i in range(heading_idx + 1, search_limit):
                p = all_paras[i]
                if p._element.findall(f".//{_DRAWING_TAG}"):
                    placeholder_para = p
                    logger.debug(
                        "doc_builder.diagram_placeholder_found role=%s para_idx=%d", role, i
                    )
                    break

            if placeholder_para is None:
                logger.info(
                    "doc_builder.diagram_no_placeholder role=%s — inserting after heading", role
                )
                new_p_elem = OxmlElement("w:p")
                all_paras[heading_idx]._element.addnext(new_p_elem)
                placeholder_para = next(
                    (p for p in doc.paragraphs if p._element is new_p_elem), None
                )
                if placeholder_para is None:
                    logger.warning(
                        "doc_builder.diagram_para_create_failed role=%s", role
                    )
                    continue

            # ── Embed first image in the placeholder slot ──────────────
            first_name, first_bytes = images[0]
            p_elem = placeholder_para._element
            pPr = p_elem.find(qn("w:pPr"))
            for child in list(p_elem):
                if child is not pPr:
                    p_elem.remove(child)
            try:
                run = placeholder_para.add_run()
                run.add_picture(BytesIO(first_bytes), width=Inches(5.5))
                logger.info(
                    "doc_builder.diagram_image_inserted role=%s file=%s bytes=%d",
                    role, first_name, len(first_bytes),
                )
            except Exception:
                logger.exception(
                    "doc_builder.diagram_image_insert_failed role=%s file=%s", role, first_name
                )

            # ── Caption for first image (when multiple images are present) ──
            anchor_elem = placeholder_para._element
            if len(images) > 1:
                role_label = "Current State" if role == "current" else "Target"
                stem = first_name.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
                caption_text = f"Figure 1: {role_label} Architecture – {stem}"
                caption_elem = _build_para_elem(caption_text)
                anchor_elem.addnext(caption_elem)
                anchor_elem = caption_elem

            # ── Additional images with captions ────────────────────────
            for idx, (fname, fbytes) in enumerate(images[1:], start=2):
                try:
                    new_p_elem = OxmlElement("w:p")
                    anchor_elem.addnext(new_p_elem)
                    img_para = next(
                        (p for p in doc.paragraphs if p._element is new_p_elem), None
                    )
                    if img_para is None:
                        continue
                    img_run = img_para.add_run()
                    img_run.add_picture(BytesIO(fbytes), width=Inches(5.5))
                    logger.info(
                        "doc_builder.diagram_image_inserted role=%s file=%s bytes=%d",
                        role, fname, len(fbytes),
                    )
                    # Caption below each additional image
                    role_label = "Current State" if role == "current" else "Target"
                    stem = fname.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
                    caption_text = f"Figure {idx}: {role_label} Architecture – {stem}"
                    caption_elem = _build_para_elem(caption_text)
                    img_para._element.addnext(caption_elem)
                    anchor_elem = caption_elem
                except Exception:
                    logger.exception(
                        "doc_builder.diagram_image_insert_failed role=%s file=%s", role, fname
                    )

    # ------------------------------------------------------------------
    # Page break before every Heading 1
    # ------------------------------------------------------------------

    @staticmethod
    def _apply_heading1_page_breaks(doc: Document) -> None:
        """Add a pageBreakBefore property to every Heading 1 except the first.

        This ensures each major section starts on a fresh page without the engineer
        having to manually insert page breaks after every generation run.
        """
        first_h1_seen = False
        for para in doc.paragraphs:
            if not (para.style.name == "Heading 1" or
                    (para.style.name.startswith("Heading") and
                     DocumentBuilder._get_heading_level(para) == 1)):
                continue
            if not first_h1_seen:
                first_h1_seen = True
                continue  # don't add a break before the very first H1
            p_elem = para._element
            pPr = p_elem.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p_elem.insert(0, pPr)
            if pPr.find(qn("w:pageBreakBefore")) is None:
                pbr = OxmlElement("w:pageBreakBefore")
                pPr.append(pbr)
        logger.info("doc_builder.heading1_page_breaks_applied")

    @staticmethod
    def _suppress_blank_heading_page_breaks(doc: Document) -> None:
        """Remove spurious page breaks that produce blank pages.

        Two passes — both operate only on paragraphs whose visible text is
        empty or whitespace-only (so content paragraphs are never touched):

        **Pass 1 — ``<w:pageBreakBefore/>`` on empty headings.**
        ``_apply_heading1_page_breaks`` adds this element to every H1 except
        the first.  Empty H1 paragraphs left in the template trigger a blank
        page because Word honours the page-break even though nothing is
        rendered.

        **Pass 2 — explicit ``<w:br w:type="page"/>`` runs in empty paragraphs.**
        Word / template editors sometimes insert a lone page-break run inside
        an empty Normal paragraph immediately following an empty heading,
        compounding the blank-page problem.  Removing the ``<w:br>`` element
        (and the now-empty run that contained it) eliminates the extra page.

        IMPORTANT: paragraphs that contain inline images have ``para.text == ""``
        but their runs hold ``<w:drawing>`` elements.  We must never touch those
        runs — only runs that are truly devoid of any visible content.
        """
        # Tags that represent real content inside a run.  A run carrying any of
        # these must NOT be deleted, even if it has no plain text.
        _CONTENT_RUN_TAGS = {
            qn("w:t"),        # text
            qn("w:br"),       # line / page break
            qn("w:drawing"),  # inline image / shape
            qn("w:pict"),     # legacy VML picture
            qn("w:object"),   # embedded OLE object
            qn("w:sym"),      # symbol glyph
            qn("w:tab"),      # tab character
        }

        removed_pbr = 0  # pageBreakBefore on empty headings
        removed_brk = 0  # explicit page-break runs in empty paragraphs

        for para in doc.paragraphs:
            if para.text.strip():
                continue  # paragraph has visible text — leave page breaks alone

            # Skip paragraphs that contain images / drawings (para.text is "" for those)
            p_elem = para._element
            if p_elem.findall(f".//{qn('w:drawing')}") or p_elem.findall(f".//{qn('w:pict')}"):
                continue

            # Pass 1: pageBreakBefore on empty headings
            if DocumentBuilder._get_heading_level(para) is not None:
                pPr = p_elem.find(qn("w:pPr"))
                if pPr is not None:
                    pbr = pPr.find(qn("w:pageBreakBefore"))
                    if pbr is not None:
                        pPr.remove(pbr)
                        removed_pbr += 1

            # Pass 2: explicit page-break runs in empty paragraphs
            for r_elem in list(p_elem.findall(qn("w:r"))):
                for br_elem in list(r_elem.findall(qn("w:br"))):
                    if br_elem.get(qn("w:type")) == "page":
                        r_elem.remove(br_elem)
                        removed_brk += 1
                # Drop the run element only if it has no meaningful content at all
                has_content = any(child.tag in _CONTENT_RUN_TAGS for child in r_elem)
                if not has_content:
                    p_elem.remove(r_elem)

        if removed_pbr:
            logger.info("doc_builder.blank_heading_page_breaks_suppressed count=%d", removed_pbr)
        if removed_brk:
            logger.info("doc_builder.blank_para_page_breaks_suppressed count=%d", removed_brk)

    # ------------------------------------------------------------------
    # SCOPE colored-box content injection
    # ------------------------------------------------------------------

    # Mapping from LLM section label keywords (lower-case, partial match) to
    # the zero-based index of the scope colored box it belongs to.
    _SCOPE_LABEL_TO_BOX: list[tuple[str, int]] = [
        ("initial understanding", 0),
        ("initial scope",         0),
        ("customer desired outcome", 1),
        ("desired outcome (customer)", 1),
        ("customer perspective",  1),
        ("desired outcome of",    1),
        ("agreed outcome",        2),
        ("desired outcome (agreed)", 2),
        ("jointly agreed",        2),
        ("oracle",                2),   # fallback for "agreed with Oracle" phrases
    ]

    def _inject_scope_boxes(
        self,
        doc: Document,
        heading_idx: int,
        next_heading_idx: int,
        content: str,
    ) -> None:
        """Handle SCOPE section injection.

        The template has three colored box *header* paragraphs (pBdr + shd)
        each followed by one or more plain *content* paragraphs (no border).
        The correct visual layout is:

            ┌─────────────────────────────────────┐
            │  Header label (colored background)  │  ← template paragraph, untouched
            └─────────────────────────────────────┘
            Body text, bullets, numbered lists…      ← plain paragraphs, LLM content

        The old approach appended content into the header paragraph via soft
        breaks, collapsing header + body into one bordered box.  This rewrite
        leaves every header paragraph untouched and injects LLM text into the
        plain content paragraphs that already exist in the template below each
        header.

        Algorithm
        ---------
        1. Walk paragraphs between the Scope heading and the next heading.
        2. Build *box_groups*: list of (header_para, [plain_content_paras]).
           A header paragraph has **both** ``<w:pBdr>`` and ``<w:shd>``.
           All subsequent paragraphs without ``<w:pBdr>`` belong to that group.
        3. Remove LLM-injected paragraphs that appeared *before* the first box.
        4. Parse the LLM ``content`` into up to three labeled blocks using
           :data:`_SCOPE_LABEL_TO_BOX`.
        5. For each group: clear the placeholder text from the first plain
           content paragraph and write the LLM lines there.  Extra lines are
           inserted as new plain paragraphs immediately after (cloning the
           paragraph properties of the template content paragraph).
        """
        import copy

        paragraphs = list(doc.paragraphs)
        body = doc.element.body
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        # ── Step 1: classify paragraphs into box groups ────────────────────
        box_groups: list[tuple] = []          # (header_para, [content_paras])
        pre_box_llm_paras: list = []
        first_box_found = False
        current_header = None
        current_content_paras: list = []

        for i in range(heading_idx + 1, next_heading_idx):
            para = paragraphs[i]
            p    = para._element
            has_pBdr = p.find(f".//{qn('w:pBdr')}") is not None
            has_shd  = p.find(f".//{qn('w:shd')}")  is not None
            is_box_header = has_pBdr and has_shd

            if is_box_header:
                if current_header is not None:
                    box_groups.append((current_header, current_content_paras))
                current_header = para
                current_content_paras = []
                first_box_found = True
            elif current_header is not None:
                current_content_paras.append(para)
            elif not first_box_found and para.text.strip():
                has_rsid = p.get(f"{{{W}}}rsidR") is not None
                if not has_rsid:
                    pre_box_llm_paras.append(para)

        if current_header is not None:
            box_groups.append((current_header, current_content_paras))

        # ── Step 2: remove stale LLM intro paragraphs ─────────────────────
        for para in pre_box_llm_paras:
            try:
                body.remove(para._element)
                logger.debug("doc_builder.scope_intro_removed text=%r", para.text[:60])
            except Exception:
                pass

        if not box_groups:
            logger.warning("doc_builder.scope_no_boxes_found — falling back to append")
            anchor = paragraphs[heading_idx]._element
            self._inject_blocks_after_element(anchor, content)
            return

        # ── Step 3: parse LLM output into 3 labeled blocks ────────────────
        blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
        box_contents: list[str] = ["", "", ""]
        current_box = 0

        for block in blocks:
            lines     = block.split("\n", 1)
            label     = lines[0].strip().rstrip(":").lower()
            body_text = lines[1].strip() if len(lines) > 1 else ""

            matched = False
            for keyword, idx in self._SCOPE_LABEL_TO_BOX:
                if keyword in label:
                    current_box   = idx
                    box_contents[idx] = body_text or block
                    matched       = True
                    break
            if not matched:
                if box_contents[current_box]:
                    box_contents[current_box] += "\n" + block
                else:
                    box_contents[current_box] = block

        # ── Step 4: inject into plain content paragraphs below each header ─
        for box_idx, (header_para, content_paras) in enumerate(box_groups):
            if box_idx >= len(box_contents):
                break
            box_text = box_contents[box_idx].strip()
            if not box_text:
                continue

            # Find the first plain paragraph (no pBdr) to use as anchor.
            # Blank separator paragraphs are skipped in favour of the first
            # paragraph that either has text or follows the header directly.
            plain_paras = [
                cp for cp in content_paras
                if cp._element.find(f".//{qn('w:pBdr')}") is None
            ]

            if plain_paras:
                target_elem = plain_paras[0]._element
                # Clear all existing runs / fldSimple (placeholder text)
                for child_tag in (qn("w:r"), qn("w:fldSimple"), qn("w:hyperlink")):
                    for ch in list(target_elem.findall(child_tag)):
                        target_elem.remove(ch)
            else:
                # No content paragraph in template — insert a new one after header
                target_elem = OxmlElement("w:p")
                header_para._element.addnext(target_elem)

            # Paragraph properties to clone for any additional new paragraphs
            pPr_template = target_elem.find(qn("w:pPr"))

            # Split content into non-empty lines
            content_lines = [ln.strip() for ln in box_text.split("\n") if ln.strip()]
            if not content_lines:
                continue

            def _make_run(text: str) -> "OxmlElement":
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                clr = OxmlElement("w:color")
                clr.set(qn("w:val"), "000000")
                rPr.append(clr)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = text
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r.append(t)
                return r

            # First line → inject into the cleared template paragraph
            target_elem.append(_make_run(content_lines[0]))

            # Remaining lines → new plain paragraphs inserted after target
            insert_after = target_elem
            for line in content_lines[1:]:
                new_p = OxmlElement("w:p")
                if pPr_template is not None:
                    new_p.append(copy.deepcopy(pPr_template))
                new_p.append(_make_run(line))
                insert_after.addnext(new_p)
                insert_after = new_p

        logger.info(
            "doc_builder.scope_boxes_filled boxes=%d removed_intro=%d",
            len(box_groups),
            len(pre_box_llm_paras),
        )

    # ------------------------------------------------------------------
    # Formatted block injection (bold labels + bullet sentences)
    # ------------------------------------------------------------------

    @staticmethod
    def _inject_formatted_blocks_after_element(anchor_elem, content: str) -> None:
        """Insert content using bold sub-topic labels and bullet-point sentences.

        Used for ARCHITECTURE COMPONENTS, IMPLEMENTATION DETAILS, and ARCHITECT
        REVIEW.  The LLM output for these sections looks like::

            Networking\\nDetails about the network...\\n\\nSecurity\\nDetails...

        Each block whose first line is a known sub-topic label (Networking, Security,
        Compute, Storage and Databases, DevOps and Management, Generation Quality,
        Data Gaps, Next Steps, Recommendations) is formatted as:

        - A bold paragraph for the label
        - One bullet paragraph per item (prefix ``– ``)

        Body items are split by newlines when the body is multi-line (e.g. bullet
        lists, numbered items from Architect Review).  For single-line dense prose
        the sentence-boundary splitter is used instead.

        Other blocks (e.g. introductory paragraphs) are inserted as plain paragraphs.

        Inserts in reverse order so that block[0] ends up immediately after anchor_elem.
        """
        blocks = [b.strip() for b in content.split("\n\n") if b.strip()]

        for block in reversed(blocks):
            lines = block.split("\n", 1)
            first_line = lines[0].strip().rstrip(":")
            body = lines[1].strip() if len(lines) > 1 else ""
            is_label = first_line.lower() in _SUB_TOPIC_LABELS

            if is_label and body:
                # Split body into individual items.  When the body is multi-line
                # (bullet lists, numbered steps) split on newlines so each item
                # gets its own paragraph.  For dense single-line prose, fall back
                # to the sentence-boundary splitter.
                if "\n" in body:
                    sentences = [s.strip() for s in body.split("\n") if s.strip()]
                else:
                    sentences = [s.strip() for s in _SENTENCE_SPLIT_RE.split(body) if s.strip()]
                if not sentences:
                    sentences = [body]
                # Insert bullet paragraphs in reverse so first ends up first.
                # Each line is built with _build_para_elem so PENDING TO REVIEW
                # markers are rendered in red.
                for sentence in reversed(sentences):
                    is_numbered = (
                        bool(sentence)
                        and sentence[0].isdigit()
                        and len(sentence) > 1
                        and sentence[1] in ".)"
                    )
                    if is_numbered:
                        anchor_elem.addnext(
                            _build_word_style_para_elem(sentence, STYLE_LIST_NUMBERED)
                        )
                    else:
                        anchor_elem.addnext(
                            _build_word_style_para_elem(
                                _strip_bullet_prefix(sentence), STYLE_LIST_BULLET
                            )
                        )
                # Insert Heading3 label before the bullets
                anchor_elem.addnext(_build_word_style_para_elem(first_line, STYLE_HEADING3))
            else:
                anchor_elem.addnext(_build_para_elem(block))

    @staticmethod
    def _inject_hierarchical_bullets_after_element(anchor_elem, content: str) -> None:
        """Inject hierarchical bullet content as separate indented paragraphs.

        Splits *content* on newlines.  Each non-empty line becomes its own
        ``<w:p>`` element inserted after *anchor_elem*.  Leading whitespace
        determines the indentation level:

        * 0 leading spaces  → L1 (``left=0`` twips)
        * 2 leading spaces  → L2 (``left=360`` twips ≈ quarter-inch)
        * 4+ leading spaces → L3 (``left=720`` twips ≈ half-inch)

        ``PENDING TO REVIEW`` markers are rendered in red.  Lines are inserted
        in reverse order so that line[0] ends up immediately after *anchor_elem*.
        """
        all_lines = [line for line in content.split("\n") if line.strip()]
        for line in reversed(all_lines):
            stripped = line.lstrip()
            indent = len(line) - len(stripped)
            if indent >= 4:
                style_id = STYLE_LIST_BULLET3
            elif indent >= 2:
                style_id = STYLE_LIST_BULLET2
            else:
                style_id = STYLE_LIST_BULLET
            anchor_elem.addnext(
                _build_word_style_para_elem(_strip_bullet_prefix(stripped), style_id)
            )

    def _append_milestone_rows(
        self, doc: Document, milestones: list[dict]
    ) -> bool:
        """Append milestone rows to the existing template milestone table.

        Locates the template table by matching the header row pattern
        ``Milestone | Target Date | Completed | Comments`` and appends one row
        per milestone entry.  ``PENDING TO REVIEW`` strings in Target Date or
        Completed cells are rendered in red.

        Returns True if rows were appended, False if the table was not found.
        """
        # Locate the milestone table by header pattern
        target_table = None
        for table in doc.tables:
            if not table.rows:
                continue
            hdrs = [
                cell.text.strip().lower()
                for cell in table.rows[0].cells
            ]
            if "milestone" in hdrs[0] and len(hdrs) >= 4:
                target_table = table
                break

        if target_table is None:
            logger.warning("doc_builder.milestone_table_not_found — cannot append rows")
            return False

        for ms in milestones:
            row = target_table.add_row()
            n_cols = len(row.cells)

            # col 0: Milestone name
            row.cells[0].paragraphs[0].add_run(ms.get("milestone", ""))

            # col 1: Target Date — PENDING TO REVIEW rendered in red
            if n_cols > 1:
                self._set_cell_with_pending_red(
                    row.cells[1], ms.get("target_date", _PENDING_MARKER)
                )

            # col 2: Completed
            if n_cols > 2:
                self._set_cell_with_pending_red(
                    row.cells[2], ms.get("completed", "")
                )

            # col 3: Comments
            if n_cols > 3:
                comments = ms.get("comments", "")
                if comments:
                    row.cells[3].paragraphs[0].add_run(comments)

        logger.info(
            "doc_builder.milestone_rows_appended count=%d", len(milestones)
        )
        return True

    @staticmethod
    def _set_cell_with_pending_red(cell, text: str) -> None:
        """Write *text* into *cell*, colouring ``PENDING TO REVIEW`` spans red."""
        para = cell.paragraphs[0]
        if _PENDING_MARKER in text:
            parts = text.split(_PENDING_MARKER)
            for p_idx, part in enumerate(parts):
                if part:
                    para.add_run(part)
                if p_idx < len(parts) - 1:
                    red_run = para.add_run(_PENDING_MARKER)
                    red_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        elif text:
            para.add_run(text)

    def _inject_structured_section(
        self,
        anchor_elem,
        data: dict,
        section_name: str,
        doc: Document | None = None,
    ) -> None:
        """Render a parsed structured-JSON section into the document.

        Called by :meth:`_inject_labeled_content` when the LLM content is
        valid JSON for a section in :data:`_STRUCTURED_OUTPUT_SECTIONS`.

        **MILESTONE PLAN** — appends rows to the existing template milestone
        table (Milestone | Target Date | Completed | Comments).  Requires
        *doc* so the table can be located via the python-docx table list.

        **HA / BACKUP / DR** — rendered as bold sub-topic labels followed by
        ``ListParagraph``-styled bullet lines.  Numbered recovery steps (lines
        whose first character is a digit) are injected without an extra ``– ``
        prefix so they render as "1. Step …" rather than "– 1. Step …".

        All items are inserted in reverse order with ``addnext`` so that the
        first logical block ends up immediately after *anchor_elem*.
        """
        upper = section_name.upper()

        if upper == "MILESTONE PLAN":
            milestones = data.get("milestones", [])
            if doc is not None and milestones:
                self._append_milestone_rows(doc, milestones)
            elif milestones:
                # Fallback: doc not available — degrade to labelled bullet blocks
                for ms in reversed(milestones):
                    name = ms.get("milestone", "")
                    target = ms.get("target_date", _PENDING_MARKER)
                    anchor_elem.addnext(
                        _build_word_style_para_elem(f"{name} — {target}", STYLE_LIST_BULLET)
                    )
        else:
            labels = _STRUCTURED_SECTION_LABELS.get(upper, [])
            for field_name, display_label in reversed(labels):
                bullets = data.get(field_name, [])
                for bullet in reversed(bullets):
                    is_numbered = (
                        bool(bullet)
                        and bullet[0].isdigit()
                        and len(bullet) > 1
                        and bullet[1] in ".)"
                    )
                    if is_numbered:
                        anchor_elem.addnext(
                            _build_word_style_para_elem(bullet, STYLE_LIST_NUMBERED)
                        )
                    else:
                        anchor_elem.addnext(
                            _build_word_style_para_elem(
                                _strip_bullet_prefix(bullet), STYLE_LIST_BULLET
                            )
                        )
                anchor_elem.addnext(_build_word_style_para_elem(display_label, STYLE_HEADING3))

    def _inject_labeled_content(
        self,
        anchor_elem,
        content: str,
        section_name: str,
        doc: Document | None = None,
    ) -> None:
        """Route to structured JSON renderer or fall back to text renderer.

        For sections in :data:`_STRUCTURED_OUTPUT_SECTIONS`:

        1. Try to parse *content* as JSON.
        2. On success → call :meth:`_inject_structured_section`.
        3. On failure (malformed JSON or unexpected payload) → log a warning
           and fall back to :meth:`_inject_formatted_blocks_after_element`.

        *doc* is forwarded to :meth:`_inject_structured_section` so the
        MILESTONE PLAN renderer can create a proper table via the python-docx
        API.  All other call sites may omit it safely (the table renderer
        degrades to a labelled-bullet fallback when ``doc`` is ``None``).

        For all other labeled-format sections, delegates directly to the
        text renderer without attempting JSON parsing.
        """
        upper = section_name.upper()
        if upper in _STRUCTURED_OUTPUT_SECTIONS:
            try:
                # Strip stray markdown code-fence wrappers before parsing.
                raw = content.strip()
                raw = re.sub(r"^```(?:json)?\s*\n?", "", raw)
                raw = re.sub(r"\n?```\s*$", "", raw)
                data = _json.loads(raw.strip())
                if not isinstance(data, dict):
                    raise ValueError(
                        f"expected JSON object, got {type(data).__name__}"
                    )
                self._inject_structured_section(anchor_elem, data, section_name, doc=doc)
                logger.debug(
                    "doc_builder.structured_injection_ok section=%s", section_name
                )
                return
            except (ValueError, _json.JSONDecodeError) as exc:
                logger.warning(
                    "doc_builder.structured_parse_failed section=%s error=%s"
                    " — falling back to text renderer",
                    section_name,
                    exc,
                )
        # Non-structured labeled section or JSON parse failure
        self._inject_formatted_blocks_after_element(anchor_elem, content)

    # ------------------------------------------------------------------
    # Table data filling
    # ------------------------------------------------------------------

    # ── Acceptance-criteria builder ──────────────────────────────────
    # Derives project-specific criteria from the architecture_analysis
    # target components.  Each entry maps a capability/metric category
    # to keywords that trigger it (checked against the flattened target
    # component list) plus a template description using ``{services}``.
    #
    # If no target components match a category it is skipped.  A generic
    # "deployment completeness" row is always emitted as the first entry.
    _CRITERIA_TEMPLATES: list[tuple[str, list[str], str]] = [
        (
            "High Availability & Resilience",
            ["data guard", "dataguard", "dns failover", "failover", "load balancer",
             "rac", "always on", "replication", "standby"],
            "The system maintains uptime and recovers automatically under "
            "simulated failure conditions using {services}.",
        ),
        (
            "Performance Validation",
            ["compute", "vm", "oke", "kubernetes", "aks", "eks", "gke",
             "instance", "autoscal"],
            "Application performance meets or exceeds the pre-migration "
            "baseline under production-representative load on {services}.",
        ),
        (
            "Data Integrity & Migration",
            ["db system", "database", "exadata", "mysql", "postgresql",
             "autonomous", "migration", "data guard", "dataguard"],
            "All data migrated completely and accurately; {services} "
            "operational with zero data loss validated.",
        ),
        (
            "Network & Connectivity",
            ["vcn", "subnet", "internet gateway", "nat gateway",
             "service gateway", "fastconnect", "vpn", "load balancer",
             "peering", "dns"],
            "Network connectivity validated end-to-end; {services} "
            "configured and tested for the target topology.",
        ),
        (
            "Storage & Backup",
            ["file storage", "object storage", "block volume", "backup",
             "replication", "block storage"],
            "Storage services operational with replication and backup "
            "validated; {services} meet retention requirements.",
        ),
        (
            "Security Posture",
            ["waf", "vault", "security", "iam", "bastion", "firewall",
             "encryption", "ssl", "tls", "certificate"],
            "Security controls verified: {services} configured and "
            "no critical/high findings in post-deployment review.",
        ),
    ]

    @classmethod
    def _build_acceptance_criteria(
        cls, architecture_analysis: dict,
    ) -> list[tuple[str, str]]:
        """Build project-specific acceptance criteria from architecture analysis.

        Scans the target-architecture component list for keywords that match
        each criteria category.  Matched services are interpolated into the
        description template.  A generic "Deployment completeness" row is
        always included first.

        Returns a list of (capability, description) tuples — max 5 entries
        to fit the template table rows.
        """
        # Flatten target component names — keep originals for display,
        # lowercased copies for keyword matching.
        target = architecture_analysis.get("target", {})
        originals: list[str] = []   # original casing for display
        lowered: list[str] = []     # lowercased for keyword matching
        if isinstance(target, dict):
            for category_items in target.values():
                if isinstance(category_items, list):
                    for item in category_items:
                        s = str(item).strip()
                        if s:
                            originals.append(s)
                            lowered.append(s.lower())
                elif isinstance(category_items, str) and category_items.strip():
                    originals.append(category_items.strip())
                    lowered.append(category_items.strip().lower())

        criteria: list[tuple[str, str]] = []

        # Always lead with deployment completeness, naming actual services
        if originals:
            # Pick up to 4 distinctive component names for the summary
            short_names: list[str] = []
            seen: set[str] = set()
            for orig in originals:
                name = orig.split("(")[0].strip()
                key = name.lower()
                if key and key not in seen and len(name) > 2:
                    seen.add(key)
                    short_names.append(name)
                if len(short_names) >= 4:
                    break
            svc_list = ", ".join(short_names) if short_names else "target OCI services"
            criteria.append((
                "Deployment Completeness",
                f"All in-scope application components deployed and operational "
                f"on OCI ({svc_list}).",
            ))
        else:
            criteria.append((
                "Deployment Completeness",
                "All in-scope application components deployed and operational "
                "on OCI as defined in the target architecture.",
            ))

        # Match each criteria template against the target components
        for capability, keywords, template in cls._CRITERIA_TEMPLATES:
            matched = [
                originals[idx].split("(")[0].strip()
                for kw in keywords
                for idx, lc in enumerate(lowered)
                if kw in lc
            ]
            # De-duplicate while preserving order
            seen_m: set[str] = set()
            unique: list[str] = []
            for m in matched:
                key_m = m.lower()
                if key_m and key_m not in seen_m:
                    seen_m.add(key_m)
                    unique.append(m)
            if unique:
                services_str = ", ".join(unique[:3])
                criteria.append((
                    capability,
                    template.format(services=services_str),
                ))
            if len(criteria) >= 5:
                break

        # Pad to at least 3 entries if we have very few components
        if len(criteria) < 3 and not originals:
            criteria.append((
                "Performance Validation",
                "Application response time remains within agreed SLA targets "
                "under production-representative load.",
            ))
            criteria.append((
                "High Availability & Resilience",
                "Failover behaviour validated: single-node failure does not "
                "cause service outage; uptime target ≥ 99.9 %.",
            ))

        return criteria[:5]

    def _fill_project_tables(
        self, doc: Document, project_context: dict | None = None
    ) -> None:
        """Fill structured table data in the DOCX template.

        Actions performed:
        - Version History: today's date when Revision Date cell is blank/placeholder.
        - Company Profile: legal name, country, industry from project context.
        - In Scope Application: application name, general description from context.
        - Application Details (General Aspects): arch type, languages, tech stack.
        - Database Tier: product, sizing, OS, scalability, availability, backup, security.
        - Application Tier: product, sizing, OS, features.
        - OCI Service Sizing / BOM: rows from inferred_metadata.oci_bom.
        - Acceptance Criteria: sample criteria for an OCI deployment project.
        - All tables: replace ``DD-MM-YYYY`` date placeholders with em-dash.
        """
        today_str = datetime.date.today().strftime("%d-%m-%Y")
        _DATE_PH = re.compile(r"^D[D\-]+M[M\-]+Y+$", re.IGNORECASE)
        ctx = project_context or {}
        meta = ctx.get("inferred_metadata") or {}

        for table in doc.tables:
            if not table.rows:
                continue

            headers = [
                cell.text.strip().lower().replace("\n", " ")
                for cell in table.rows[0].cells
            ]
            h0 = headers[0] if headers else ""

            # ── Version History ────────────────────────────────────────
            is_version_table = (
                "revision date" in headers
                or ("version #" in headers and "revised by" in headers)
                or ("version" in headers and "revised by" in headers)
            )
            if is_version_table and len(table.rows) > 1:
                data_row = table.rows[1]
                for ci, header in enumerate(headers):
                    if ci >= len(data_row.cells):
                        break
                    cell = data_row.cells[ci]
                    cell_text = cell.text.strip()
                    if "date" in header:
                        # Always overwrite — template may ship with a real-looking
                        # date (e.g. "01-10-2025") that _DATE_PH wouldn't match.
                        self._set_cell_text(cell, today_str)
                        logger.info("doc_builder.version_history_date_set date=%s", today_str)
                    elif "revised by" in header and not cell_text:
                        self._set_cell_text(cell, "Oracle Labs")
                        logger.info("doc_builder.version_history_author_set")
                continue  # DD-MM-YYYY pass handled above

            # ── Company Profile ────────────────────────────────────────
            # Detected by "legal name" appearing in the first column headers.
            if any("legal name" in h for h in headers):
                for row in table.rows:
                    label = row.cells[0].text.strip().lower() if row.cells else ""
                    val_cell = row.cells[1] if len(row.cells) > 1 else None
                    if val_cell is None:
                        continue
                    current = val_cell.text.strip()
                    if not current or current == " ":
                        if "legal name" in label:
                            self._set_cell_text(val_cell, self.customer_name)
                        elif "country" in label:
                            country = meta.get("country") or ""
                            if country:
                                self._set_cell_text(val_cell, country)
                        elif "industry" in label or "selling" in label:
                            industry = (
                                meta.get("industry")
                                or ctx.get("industry")
                                or ""
                            )
                            if industry:
                                self._set_cell_text(val_cell, industry)
                        elif "description" in label or "company" in label:
                            desc = meta.get("company_description") or ""
                            if desc:
                                self._set_cell_text(val_cell, desc)
                logger.info("doc_builder.company_profile_filled")
                continue

            # ── In Scope Application ───────────────────────────────────
            # Detected by first header = "application name".
            # Fills the 2-column label→value table with application info
            # drawn from inferred_metadata (not engagement scope).
            if "application name" in h0:
                for row in table.rows:
                    label = row.cells[0].text.strip().lower() if row.cells else ""
                    val_cell = row.cells[1] if len(row.cells) > 1 else None
                    if val_cell is None:
                        continue
                    current = val_cell.text.strip()
                    if "application name" in label:
                        # Always overwrite to fix the Customer1Project1 concatenation
                        app_name = self.project_name or self.customer_name or ""
                        self._set_cell_text(val_cell, app_name.strip())
                    elif "general description" in label and not current:
                        # Use app_required_features (what the app does) from
                        # inferred metadata, NOT the engagement scope.
                        desc = (
                            meta.get("app_required_features")
                            or meta.get("company_description")
                            or ""
                        )
                        if desc:
                            self._set_cell_text(val_cell, desc)
                    elif "key technologies" in label and not current:
                        techs = meta.get("used_technologies") or ""
                        if techs:
                            self._set_cell_text(val_cell, techs)
                    elif "running on" in label and not current:
                        platform = ctx.get("cloud") or ""
                        val = (
                            f"{platform} → OCI (Oracle Cloud Infrastructure)"
                            if platform and platform.upper() not in {"OCI", "ORACLE"}
                            else "OCI (Oracle Cloud Infrastructure)"
                        )
                        self._set_cell_text(val_cell, val)
                logger.info("doc_builder.in_scope_application_filled")
                continue

            # ── Application Details (General Aspects) ─────────────────
            # Detected by first header = "general aspects".
            if "general aspects" in h0:
                _APP_DETAIL_MAP = {
                    "application architecture": "app_architecture_type",
                    "development language":     "development_languages",
                    "hardware dependencies":    "hardware_dependencies",
                    "used technologies":        "used_technologies",
                }
                for row in table.rows[1:]:  # skip header row
                    if len(row.cells) < 2:
                        continue
                    label = row.cells[0].text.strip().lower()
                    val_cell = row.cells[1]
                    if val_cell.text.strip():
                        continue  # already has content
                    for key_fragment, meta_key in _APP_DETAIL_MAP.items():
                        if key_fragment in label:
                            value = meta.get(meta_key) or ""
                            if value:
                                self._set_cell_text(val_cell, value)
                            break
                logger.info("doc_builder.app_details_filled")
                continue

            # ── Database Tier ──────────────────────────────────────────
            # Detected by first header = "database tier".
            if "database tier" in h0:
                _DB_MAP = {
                    "product, edition, version":    "db_product_edition",
                    "server sizing":                "db_server_sizing",
                    "requirements":                 "db_os_requirements",
                    "size and growth":              "db_size_and_growth",
                    "scalability":                  "db_scalability",
                    "availability":                 "db_availability",
                    "backup":                       "db_backup",
                    "security":                     "db_security",
                }
                # Row 0 is merged header; Row 1 is sub-header "Relevant Aspects / Details"
                for row in table.rows[2:]:
                    if len(row.cells) < 2:
                        continue
                    label = row.cells[0].text.strip().lower()
                    val_cell = row.cells[1]
                    if val_cell.text.strip():
                        continue
                    for key_fragment, meta_key in _DB_MAP.items():
                        if key_fragment in label:
                            value = meta.get(meta_key) or ""
                            if value:
                                self._set_cell_text(val_cell, value)
                            break
                logger.info("doc_builder.database_tier_filled")
                continue

            # ── Application Tier ───────────────────────────────────────
            # Detected by first header = "application tier".
            if "application tier" in h0:
                _APP_TIER_MAP = {
                    "product, edition, version":        "app_product_version",
                    "server sizing":                    "app_server_sizing",
                    "hardware dependencies":            "hardware_dependencies",
                    "os and dependencies":              "app_os_dependencies",
                    "required functionality":           "app_required_features",
                }
                for row in table.rows[2:]:
                    if len(row.cells) < 2:
                        continue
                    label = row.cells[0].text.strip().lower()
                    val_cell = row.cells[1]
                    if val_cell.text.strip():
                        continue
                    for key_fragment, meta_key in _APP_TIER_MAP.items():
                        if key_fragment in label:
                            value = meta.get(meta_key) or ""
                            if value:
                                self._set_cell_text(val_cell, value)
                            break
                logger.info("doc_builder.app_tier_filled")
                continue

            # ── OCI Service Sizing / Bill of Materials ─────────────────
            # Detected by first header = "oci service name".
            if "oci service name" in h0:
                bom: list[dict] = meta.get("oci_bom") or []
                if bom:
                    # Overwrite placeholder rows first, then add more if needed
                    placeholder_rows = [
                        r for r in table.rows[1:]
                        if r.cells[0].text.strip().lower().startswith("service")
                    ]
                    for ri, entry in enumerate(bom):
                        if ri < len(placeholder_rows):
                            row = placeholder_rows[ri]
                        else:
                            row = table.add_row()
                        cells = row.cells
                        if len(cells) >= 1:
                            self._set_cell_text(cells[0], entry.get("service", ""))
                        if len(cells) >= 2:
                            self._set_cell_text(cells[1], entry.get("sizing_unit", ""))
                        if len(cells) >= 3:
                            self._set_cell_text(cells[2], entry.get("amount", ""))
                        if len(cells) >= 4:
                            self._set_cell_text(cells[3], entry.get("comments", ""))
                    # Blank out any remaining placeholder rows that weren't used
                    for row in placeholder_rows[len(bom):]:
                        for cell in row.cells:
                            self._set_cell_text(cell, "")
                    logger.info("doc_builder.bom_filled entries=%d", len(bom))
                # Apply visual styling regardless of whether BOM data was provided:
                # light blue header row, dark grayish-green data rows, fix header typo.
                self._style_bom_table(table)
                continue

            # ── Acceptance Criteria ────────────────────────────────────
            # Detected by "acceptance criteria" in col 1 of the header row.
            # Criteria are derived from the target architecture components
            # so they reflect the actual services being validated.
            if len(headers) > 1 and "acceptance criteria" in headers[1]:
                arch_analysis = ctx.get("architecture_analysis") or {}
                criteria = self._build_acceptance_criteria(arch_analysis)
                data_rows = table.rows[1:]  # skip header row
                for ri, (capability, description) in enumerate(criteria):
                    if ri >= len(data_rows):
                        break
                    row = data_rows[ri]
                    if len(row.cells) >= 2:
                        self._set_cell_text(row.cells[0], capability)
                        self._set_cell_text(row.cells[1], description)
                        if len(row.cells) >= 3:
                            self._set_cell_text(row.cells[2], "Pending")
                        if len(row.cells) >= 4:
                            self._set_cell_text(row.cells[3], "—")
                logger.info(
                    "doc_builder.acceptance_criteria_filled count=%d",
                    len(criteria),
                )
                continue

            # ── General DD-MM-YYYY placeholder sweep ──────────────────
            for row in table.rows[1:]:
                for cell in row.cells:
                    if _DATE_PH.match(cell.text.strip()):
                        self._set_cell_text(cell, "\u2014")  # em-dash

        logger.info("doc_builder.project_tables_filled")

    @staticmethod
    def _set_cell_text(cell, text: str) -> None:
        """Set the text of the first paragraph in a table cell, preserving formatting."""
        for para in cell.paragraphs:
            # Clear all runs
            for run in para.runs:
                run.text = ""
            # Write into the first run, or create one
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            return  # only touch the first paragraph

    @staticmethod
    def _set_cell_fill(cell, hex_color: str) -> None:
        """Set the background fill colour of a table cell.

        Args:
            cell: python-docx ``_Cell`` object.
            hex_color: Six-character hex colour string without the ``#``
                prefix, e.g. ``"BDD7EE"`` for light blue.
        """
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for existing in tcPr.findall(qn("w:shd")):
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    @staticmethod
    def _set_cell_text_color(cell, hex_color: str, bold: bool = False) -> None:
        """Apply text colour (and optionally bold) to every run in a table cell.

        Args:
            cell: python-docx ``_Cell`` object.
            hex_color: Six-character hex colour string, e.g. ``"FFFFFF"``
                for white or ``"1F3864"`` for dark navy.
            bold: When *True*, also sets ``<w:b/>`` on each run's rPr.
        """
        for para in cell.paragraphs:
            for run in para.runs:
                r_elem = run._r
                rPr = r_elem.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    r_elem.insert(0, rPr)
                for existing in rPr.findall(qn("w:color")):
                    rPr.remove(existing)
                clr = OxmlElement("w:color")
                clr.set(qn("w:val"), hex_color)
                rPr.append(clr)
                if bold and rPr.find(qn("w:b")) is None:
                    rPr.insert(0, OxmlElement("w:b"))

    @staticmethod
    def _style_bom_table(table) -> None:
        """Apply OCI BOM table visual styling — unified with the Acceptance Criteria table.

        Table-level
        -----------
        * Style ``BasicTable09Redwood`` (matches Acceptance Criteria table).
        * All four outer borders: single, sz=6, color ``312D2A``.

        Header row
        ----------
        * Fill: dark navy ``001F5B``.
        * Text: white ``FFFFFF``, bold, centred.
        * Also fixes the "Sizing Units" header text/typo and removes any
          duplicate second paragraph inside the cell.

        Data rows
        ---------
        * First column fill: ``D4DFDF`` (teal-grey, matches Acceptance Criteria).
        * Other columns fill: ``E9EFEF`` (lighter teal-grey).
        * Text: black ``000000``.
        """
        if not table.rows:
            return

        # ── Table-level style and borders ──────────────────────────────
        tbl_elem = table._tbl
        tblPr = tbl_elem.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_elem.insert(0, tblPr)

        # tblStyle
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is None:
            tblStyle = OxmlElement("w:tblStyle")
            tblPr.insert(0, tblStyle)
        tblStyle.set(qn("w:val"), "BasicTable09Redwood")

        # tblBorders — replace entirely
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "312D2A")
            b.set(qn("w:themeColor"), "text1")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # tblLook
        for old in tblPr.findall(qn("w:tblLook")):
            tblPr.remove(old)
        tblLook = OxmlElement("w:tblLook")
        tblLook.set(qn("w:val"), "04A0")
        tblLook.set(qn("w:firstRow"), "1")
        tblLook.set(qn("w:lastRow"), "0")
        tblLook.set(qn("w:firstColumn"), "1")
        tblLook.set(qn("w:lastColumn"), "0")
        tblLook.set(qn("w:noHBand"), "0")
        tblLook.set(qn("w:noVBand"), "1")
        tblPr.append(tblLook)

        # ── Header row ─────────────────────────────────────────────────
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell_text = cell.text.strip()
            if "sizing" in cell_text.lower() and "unit" in cell_text.lower():
                DocumentBuilder._set_cell_text(cell, "Sizing Units (ex. vCPUs)")
                # Remove duplicate second <w:p> that causes "…(ex. vCPUs)\n(ex. vCPUs)"
                tc = cell._tc
                for extra_p in tc.findall(qn("w:p"))[1:]:
                    tc.remove(extra_p)
            DocumentBuilder._set_cell_fill(cell, "001F5B")
            DocumentBuilder._set_cell_text_color(cell, "FFFFFF", bold=True)
            # Centre-align header text
            for para in cell.paragraphs:
                pPr = para._p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    para._p.insert(0, pPr)
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    pPr.append(jc)
                jc.set(qn("w:val"), "center")

        # ── Data rows ──────────────────────────────────────────────────
        for row in table.rows[1:]:
            for ci, cell in enumerate(row.cells):
                fill = "D4DFDF" if ci == 0 else "E9EFEF"
                DocumentBuilder._set_cell_fill(cell, fill)
                DocumentBuilder._set_cell_text_color(cell, "000000")

    def _delete_section_from_template(self, doc: Document, section_name: str) -> bool:
        """Remove a section's heading and all body content up to the next same-level heading.

        Used when a section is excluded from generation — ensures the template heading
        doesn't appear in the final DOCX as an empty stub.
        """
        raw_keyword = SECTION_HEADING_KEYWORDS.get(section_name.upper(), section_name.lower())
        exact_match = raw_keyword.startswith("=")
        keyword = raw_keyword[1:] if exact_match else raw_keyword

        paragraphs = list(doc.paragraphs)

        # Locate the section heading
        heading_idx: int | None = None
        heading_level = 1
        for i, para in enumerate(paragraphs):
            if not self._is_heading_style(para):
                continue
            para_text = para.text.strip().lower()
            matched = (para_text == keyword.lower()) if exact_match else (keyword.lower() in para_text)
            if matched:
                heading_idx = i
                heading_level = self._get_heading_level(para) or 1
                break

        if heading_idx is None:
            logger.debug("doc_builder.delete_section_not_found section=%s", section_name)
            return False

        # Find where this section ends (next heading at same or higher hierarchy level)
        end_idx = len(paragraphs)
        for i in range(heading_idx + 1, len(paragraphs)):
            if self._is_heading_style(paragraphs[i]):
                lvl = self._get_heading_level(paragraphs[i]) or 1
                if lvl <= heading_level:
                    end_idx = i
                    break

        # Remove all paragraphs in [heading_idx, end_idx)
        for para in paragraphs[heading_idx:end_idx]:
            p = para._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)

        logger.info(
            "doc_builder.section_deleted section=%s removed_paragraphs=%d",
            section_name,
            end_idx - heading_idx,
        )
        return True

    def _inject_section(self, doc: Document, section_name: str, content: str) -> bool:
        """Find heading in template, remove placeholder paragraphs, inject content."""
        raw_keyword = SECTION_HEADING_KEYWORDS.get(section_name.upper(), section_name.lower())

        # A keyword prefixed with "=" requires an EXACT heading-text match
        # (case-insensitive) rather than a substring search.  Used for short
        # keywords like "scope" that would otherwise match longer headings.
        exact_match = raw_keyword.startswith("=")
        keyword = raw_keyword[1:] if exact_match else raw_keyword

        paragraphs = list(doc.paragraphs)

        # Find the matching heading paragraph (body-level paragraphs only)
        heading_idx: int | None = None
        heading_level = 1
        for i, para in enumerate(paragraphs):
            if not self._is_heading_style(para):
                continue
            para_text = para.text.strip().lower()
            if exact_match:
                matched = (para_text == keyword.lower())
            else:
                matched = (keyword.lower() in para_text)
            if matched:
                heading_idx = i
                heading_level = self._get_heading_level(para) or 1
                break

        if heading_idx is None:
            # Fallback: search inside table cells — python-docx excludes these from doc.paragraphs
            table_elem = self._find_heading_in_tables(doc, keyword)
            if table_elem is not None:
                logger.info(
                    "doc_builder.heading_in_table section=%s — injecting after table element",
                    section_name,
                )
                self._inject_blocks_after_element(table_elem, content)
                content_blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
                logger.info(
                    "doc_builder.section_injected section=%s blocks=%d via table_fallback",
                    section_name,
                    len(content_blocks),
                )
                return True

            available_headings = [p.text for p in paragraphs if self._is_heading_style(p)]
            logger.warning(
                "doc_builder.heading_not_found section=%s keyword=%r — "
                "available template headings: %s",
                section_name,
                keyword,
                available_headings,
            )
            return False

        # Find range: heading+1 → next same-or-higher-level heading
        next_heading_idx = len(paragraphs)
        for i in range(heading_idx + 1, len(paragraphs)):
            lvl = self._get_heading_level(paragraphs[i])
            if lvl is not None and lvl <= heading_level:
                next_heading_idx = i
                break

        body = doc.element.body
        content_blocks = [b.strip() for b in content.split("\n\n") if b.strip()]

        # ── SCOPE special handling ───────────────────────────────────────
        # Retains the colored scope-box paragraphs from the template, removes
        # only LLM-generated intro text, and injects content inside each box.
        if section_name.upper() == "SCOPE":
            self._inject_scope_boxes(doc, heading_idx, next_heading_idx, content)
            logger.info(
                "doc_builder.section_injected section=%s blocks=%d (scope_boxes)",
                section_name, len(content_blocks),
            )
            return True

        # ── Single-sentence sections: truncate to first paragraph ────────
        # These sections must produce exactly ONE introductory sentence.
        # The LLM is instructed accordingly but may still emit extra
        # paragraphs.  Silently truncate so only the first paragraph is
        # injected, preventing unwanted text above the following table.
        if section_name.upper() in _SINGLE_SENTENCE_SECTIONS:
            first_block = content.split("\n\n")[0].strip()
            if first_block:
                content = first_block
                content_blocks = [first_block]
                logger.debug(
                    "doc_builder.content_truncated_single_sentence section=%s",
                    section_name,
                )

        # ── Full-clear sections ──────────────────────────────────────────
        # For sections that are 100% LLM-generated, remove ALL non-heading
        # template paragraphs between the section heading and the next
        # heading, then inject the LLM content immediately after the heading.
        if section_name.upper() in _FULL_CLEAR_SECTIONS:
            for i in range(heading_idx + 1, next_heading_idx):
                para = paragraphs[i]
                if self._get_heading_level(para) is None:
                    try:
                        body.remove(para._element)
                    except Exception:
                        pass
            anchor_elem = paragraphs[heading_idx]._element
            _use_formatted    = section_name.upper() in _LABELED_FORMAT_SECTIONS
            _use_hierarchical = section_name.upper() in _HIERARCHICAL_BULLET_SECTIONS
            if _use_formatted:
                self._inject_labeled_content(anchor_elem, content, section_name, doc=doc)
            elif _use_hierarchical:
                self._inject_hierarchical_bullets_after_element(anchor_elem, content)
            else:
                self._inject_blocks_after_element(anchor_elem, content)
            logger.info(
                "doc_builder.section_injected section=%s blocks=%d (full_clear%s%s)",
                section_name, len(content_blocks),
                ",formatted" if _use_formatted else "",
                ",hierarchical" if _use_hierarchical else "",
            )
            return True

        # ── Normal sections ──────────────────────────────────────────────
        # Remove placeholder paragraphs (but not sub-headings or real content)
        for i in range(heading_idx + 1, next_heading_idx):
            para = paragraphs[i]
            if _PLACEHOLDER_RE.search(para.text):
                body.remove(para._element)

        # Determine the injection anchor.
        # Advance the anchor past any template intro text so LLM content is
        # appended after it (never before surviving template paragraphs).
        anchor_elem = paragraphs[heading_idx]._element
        for i in range(heading_idx + 1, next_heading_idx):
            para = paragraphs[i]
            lvl = self._get_heading_level(para)
            if lvl is not None:
                break  # hit a sub-heading — keep current anchor
            if para.text.strip() and not _PLACEHOLDER_RE.search(para.text):
                anchor_elem = para._element  # advance anchor past intro para(s)

        if section_name.upper() in _LABELED_FORMAT_SECTIONS:
            self._inject_labeled_content(anchor_elem, content, section_name, doc=doc)
        else:
            self._inject_blocks_after_element(anchor_elem, content)
        logger.info("doc_builder.section_injected section=%s blocks=%d", section_name, len(content_blocks))
        return True

    def _inject_after_known_section(
        self,
        doc: Document,
        section_name: str,
        content: str,
        after_keyword: str,
    ) -> bool:
        """Insert a brand-new heading + content immediately after a sibling section.

        Used for sections (e.g. BACKUP STRATEGY, DISASTER RECOVERY) that do not
        have their own heading in the DOCX template but should appear right after
        a section that does (e.g. HIGH AVAILABILITY).

        The new heading paragraph is built by cloning the paragraph-properties
        (``<w:pPr>``) of the anchor heading so that template styles (font, colour,
        spacing, page-break-before) are preserved, while the text run is created
        fresh.  Returns True on success, False when the anchor cannot be found.
        """
        import copy
        from lxml import etree

        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paragraphs = list(doc.paragraphs)

        # Locate the anchor heading.
        anchor_idx: int | None = None
        anchor_level = 1
        for i, para in enumerate(paragraphs):
            if self._is_heading_style(para) and after_keyword.lower() in para.text.lower():
                anchor_idx = i
                anchor_level = self._get_heading_level(para) or 1
                break

        if anchor_idx is None:
            logger.warning(
                "doc_builder.inject_after_anchor_not_found section=%s anchor_keyword=%r",
                section_name,
                after_keyword,
            )
            return False

        # Find the end of the anchor section (first same-or-higher heading after it).
        end_idx = len(paragraphs)
        for i in range(anchor_idx + 1, len(paragraphs)):
            lvl = self._get_heading_level(paragraphs[i])
            if lvl is not None and lvl <= anchor_level:
                end_idx = i
                break

        # Insert after the last paragraph of the anchor section.
        insert_after_elem = (
            paragraphs[end_idx - 1]._element
            if end_idx > anchor_idx + 1
            else paragraphs[anchor_idx]._element
        )

        # Build a new heading element: clone pPr from anchor (preserves style),
        # attach a fresh text run.
        source_pPr = paragraphs[anchor_idx]._element.find(f"{{{ns}}}pPr")
        heading_elem = etree.Element(f"{{{ns}}}p")
        if source_pPr is not None:
            heading_elem.append(copy.deepcopy(source_pPr))
        r_elem = etree.SubElement(heading_elem, f"{{{ns}}}r")
        t_elem = etree.SubElement(r_elem, f"{{{ns}}}t")
        t_elem.text = section_name.title()

        insert_after_elem.addnext(heading_elem)

        # Inject body content after the new heading.
        _use_formatted = section_name.upper() in _LABELED_FORMAT_SECTIONS
        _use_hierarchical = section_name.upper() in _HIERARCHICAL_BULLET_SECTIONS
        if _use_formatted:
            self._inject_labeled_content(heading_elem, content, section_name, doc=doc)
        elif _use_hierarchical:
            self._inject_hierarchical_bullets_after_element(heading_elem, content)
        else:
            self._inject_blocks_after_element(heading_elem, content)

        content_blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
        logger.info(
            "doc_builder.section_injected_after_anchor section=%s anchor=%r blocks=%d (full_clear%s%s)",
            section_name,
            after_keyword,
            len(content_blocks),
            ",formatted" if _use_formatted else "",
            ",hierarchical" if _use_hierarchical else "",
        )
        return True

    def _append_section(self, doc: Document, section_name: str, content: str) -> None:
        """Fallback: append section at end of document.

        For sections in :data:`_LABELED_FORMAT_SECTIONS` (e.g. ARCHITECT REVIEW),
        the formatted bold-label + bullet injection is used so that sub-topic
        structure is preserved even when the section is not found in the template.
        """
        heading = doc.add_heading(section_name.title(), level=1)
        if section_name.upper() in _LABELED_FORMAT_SECTIONS:
            self._inject_labeled_content(heading._element, content, section_name, doc=doc)
        else:
            for block in content.split("\n\n"):
                if block.strip():
                    doc.add_paragraph(block.strip())

    def _find_heading_in_tables(self, doc: Document, keyword: str):
        """Search table cells for a heading paragraph matching keyword.

        Returns the parent <w:tbl> XML element so content can be inserted after
        the entire table, or None if not found.
        """
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._is_heading_style(para) and keyword.lower() in para.text.lower():
                            # para._element → <w:p>  parent → <w:tc>  parent → <w:tr>  parent → <w:tbl>
                            return para._element.getparent().getparent().getparent()
        return None

    @staticmethod
    def _inject_blocks_after_element(anchor_elem, content: str) -> None:
        """Insert content paragraphs as next siblings of anchor_elem.

        Inserts in reverse order so that block[0] ends up immediately after
        anchor_elem.  Uses :func:`_build_para_elem` so that any
        ``PENDING TO REVIEW`` markers within the text are rendered in red.
        """
        content_blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
        for block in reversed(content_blocks):
            anchor_elem.addnext(_build_para_elem(block))

    @staticmethod
    def _is_heading_style(para) -> bool:
        """Return True if the paragraph uses any Heading-based style (including custom derived styles)."""
        style = para.style
        while style is not None:
            if style.name.startswith("Heading"):
                return True
            style = getattr(style, "base_style", None)
        return False

    @staticmethod
    def _get_heading_level(para) -> int | None:
        style_name = para.style.name
        if style_name.startswith("Heading"):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 1
        return None
